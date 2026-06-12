# scripts/generate_dexuatcapnhat.py

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_heading_styled(doc, text, level, color=RGBColor(0, 51, 102)):
    h = doc.add_heading(text, level=level)
    run = h.runs[0]
    run.font.name = 'Times New Roman'
    run.font.color.rgb = color
    run.bold = True
    if level == 1:
        run.font.size = Pt(14)
        h.paragraph_format.space_before = Pt(16)
        h.paragraph_format.space_after = Pt(6)
    elif level == 2:
        run.font.size = Pt(12)
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(4)
    elif level == 3:
        run.font.size = Pt(11.5)
        h.paragraph_format.space_before = Pt(8)
        h.paragraph_format.space_after = Pt(3)
    return h

def add_paragraph_styled(doc, text="", bold_prefix=""):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    if bold_prefix:
        r_bold = p.add_run(bold_prefix)
        r_bold.bold = True
        r_bold.font.name = 'Times New Roman'
        r_bold.font.size = Pt(11)
        
    r_text = p.add_run(text)
    r_text.font.name = 'Times New Roman'
    r_text.font.size = Pt(11)
    return p

def main():
    doc_path = "report/dexuatcapnhat.docx"
    doc = Document()
    
    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Document Header Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = title.add_run("ĐỀ XUẤT CẬP NHẬT VÀ BỔ SUNG CHI TIẾT BÁO CÁO ĐỒ ÁN IE212\n(ĐỐI CHIẾU VÀ SỬA ĐỔI THEO Ý KIẾN GÓP Ý)")
    run_title.font.name = 'Times New Roman'
    run_title.font.size = Pt(14)
    run_title.bold = True
    run_title.font.color.rgb = RGBColor(0, 51, 102)
    title.paragraph_format.space_after = Pt(18)

    # I. CHỈNH SỬA SỐ LIỆU CHO KHỚP NOTEBOOK
    add_heading_styled(doc, "I. CHỈNH SỬA SỐ LIỆU ĐỂ ĐẢM BẢO KHỚP 100% VỚI NOTEBOOK", level=1)
    
    add_paragraph_styled(doc, 
        "Vị trí thay thế: Bảng 1 (Kết quả đánh giá mô hình trong Môi trường Thực nghiệm tĩnh trên thang MinMaxScaler). "
        "Cần cập nhật lại các sai số nhỏ sau để khớp tuyệt đối với CELL 18 của notebook:")
    
    # Bảng 1 update values
    t1 = doc.add_table(rows=1, cols=4)
    t1.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr1 = t1.rows[0].cells
    cols1 = ["Mô hình", "Chỉ số trong Báo cáo cũ", "Giá trị chính xác (Notebook)", "Hành động"]
    for idx, c_name in enumerate(cols1):
        hdr1[idx].text = c_name
        set_cell_background(hdr1[idx], "003366")
        set_cell_margins(hdr1[idx], top=80, bottom=80, left=80, right=80)
        run = hdr1[idx].paragraphs[0].runs[0]
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(10)
    
    t1_data = [
        ["LSTM (Static MinMaxScaler)", "MSE = 0.002080", "MSE = 0.002078", "Thay thế giá trị trong Bảng 1"],
        ["Hybrid No-Gate (Static MinMaxScaler)", "MSE = 0.002045", "MSE = 0.002046", "Thay thế giá trị trong Bảng 1"],
        ["Hybrid Graph-Gate (Static MinMaxScaler)", "MSE = 0.002045", "MSE = 0.002046", "Thay thế giá trị trong Bảng 1"]
    ]
    for r_data in t1_data:
        row_cells = t1.add_row().cells
        for c_idx, val in enumerate(r_data):
            row_cells[c_idx].text = val
            set_cell_margins(row_cells[c_idx], top=60, bottom=60, left=60, right=60)
            row_cells[c_idx].paragraphs[0].runs[0].font.name = 'Times New Roman'
            row_cells[c_idx].paragraphs[0].runs[0].font.size = Pt(10)

    doc.add_paragraph()

    add_paragraph_styled(doc, 
        "Vị trí thay thế: Bảng 2 (Kết quả đánh giá mô hình trong Môi trường Thực nghiệm tĩnh trên đơn vị giá USD thực tế). "
        "Cập nhật lại các tỷ lệ Directional Accuracy (DA) của LSTM và hai mô hình lai để khắc phục lệch số liệu nghiêm trọng:")
    
    # Bảng 2 update values
    t2 = doc.add_table(rows=1, cols=4)
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr2 = t2.rows[0].cells
    cols2 = ["Mô hình (USD thực tế)", "DA trong Báo cáo cũ", "DA chính xác (Notebook)", "Hành động"]
    for idx, c_name in enumerate(cols2):
        hdr2[idx].text = c_name
        set_cell_background(hdr2[idx], "003366")
        set_cell_margins(hdr2[idx], top=80, bottom=80, left=80, right=80)
        run = hdr2[idx].paragraphs[0].runs[0]
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(10)
        
    t2_data = [
        ["LSTM thuần (Close-only baseline)", "40.0%", "41.6%", "Sửa trong Bảng 2 và đoạn văn phân tích dưới bảng"],
        ["Hybrid LSTM-GNN No-Gate", "41.4%", "42.2%", "Sửa trong Bảng 2 và đoạn văn phân tích dưới bảng"],
        ["Hybrid LSTM-GNN Graph-Gate", "42.0%", "43.4%", "Sửa trong Bảng 2 và đoạn văn phân tích dưới bảng"]
    ]
    for r_data in t2_data:
        row_cells = t2.add_row().cells
        for c_idx, val in enumerate(r_data):
            row_cells[c_idx].text = val
            set_cell_margins(row_cells[c_idx], top=60, bottom=60, left=60, right=60)
            row_cells[c_idx].paragraphs[0].runs[0].font.name = 'Times New Roman'
            row_cells[c_idx].paragraphs[0].runs[0].font.size = Pt(10)

    doc.add_paragraph()

    # II. KHÔI PHỤC BẢNG 7 / BẢNG 6 (TÀI CHÍNH TĨNH)
    add_heading_styled(doc, "II. KHÔI PHỤC BẢNG SỐ LIỆU TÀI CHÍNH TĨNH BỊ LỖI (BẢNG 5 HOẶC 6 TRONG BÁO CÁO)", level=1)
    add_paragraph_styled(doc, 
        "Vị trí chèn: Tại phần 4.2.2 (Kết quả backtest trong Môi trường Thực nghiệm tĩnh), trước đoạn văn 'Từ kết quả Bảng 6...'. "
        "Bảng số liệu tài chính tĩnh bị mất đã được khôi phục chính xác từ notebook output như sau:")
    
    t_fin = doc.add_table(rows=1, cols=5)
    t_fin.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_f = t_fin.rows[0].cells
    cols_f = ["Chiến lược giao dịch", "Mô hình áp dụng", "Lợi nhuận tích lũy (CR)", "Hệ số Sharpe (SR)", "Max Drawdown (MDD)"]
    for idx, c_name in enumerate(cols_f):
        hdr_f[idx].text = c_name
        set_cell_background(hdr_f[idx], "003366")
        set_cell_margins(hdr_f[idx], top=80, bottom=80, left=60, right=60)
        run = hdr_f[idx].paragraphs[0].runs[0]
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(9.5)
        
    fin_data = [
        ["Long-only (Mua vị thế tăng)", "Linear Regression", "2.592%", "0.519", "-16.994%"],
        ["Long-only (Mua vị thế tăng)", "LSTM", "-17.871%", "-3.438", "-22.329%"],
        ["Long-only (Mua vị thế tăng)", "Hybrid No-Gate", "-17.103%", "-2.652", "-26.487%"],
        ["Long-only (Mua vị thế tăng)", "Hybrid Graph-Gate", "-14.525%", "-2.195", "-26.339%"],
        ["Long-Short (K=2, Cặp đối ứng)", "Linear Regression", "-11.276%", "-2.281", "-14.824%"],
        ["Long-Short (K=2, Cặp đối ứng)", "LSTM", "-37.151%", "-6.162", "-38.101%"],
        ["Long-Short (K=2, Cặp đối ứng)", "Hybrid No-Gate", "-26.808%", "-4.349", "-30.784%"],
        ["Long-Short (K=2, Cặp đối ứng)", "Hybrid Graph-Gate", "-27.190%", "-4.466", "-30.463%"],
        ["Benchmark thụ động", "Buy & Hold (Thị trường)", "1.538%", "0.387", "-17.277%"]
    ]
    for r_data in fin_data:
        row_cells = t_fin.add_row().cells
        for c_idx, val in enumerate(r_data):
            row_cells[c_idx].text = val
            set_cell_margins(row_cells[c_idx], top=50, bottom=50, left=50, right=50)
            row_cells[c_idx].paragraphs[0].runs[0].font.name = 'Times New Roman'
            row_cells[c_idx].paragraphs[0].runs[0].font.size = Pt(9)
            row_cells[c_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # III. BỔ SUNG 2 MÔ HÌNH TSN VÀ ATTENTION MỚI
    add_heading_styled(doc, "III. BỔ SUNG KẾT QUẢ VÀ THUẬT TOÁN CHO 2 MÔ HÌNH TSN MỚI", level=1)
    add_paragraph_styled(doc, 
        "Vị trí chèn: Chèn vào cuối Chương 3 (Mục 3.2.2.4 mới) và Chương 4 (Mục 4.1.1.1 và 4.1.1.2 mới). "
        "Trình bày chi tiết thuật toán làm giàu đặc trưng Xu hướng-Mùa vụ-Nhiễu (TSN), cơ chế Temporal Attention và bảng kết quả Ablation Study với 6 mô hình đầy đủ.")
    
    add_paragraph_styled(doc, "1. Công thức tính toán đặc trưng TSN động tại bước t (Chương 3):", bold_prefix="Cải tiến đặc trưng: ")
    add_paragraph_styled(doc, 
        "- Trend: $MA_{50} = \\frac{1}{50}\\sum_{i=0}^{49} Close_{t-i}$, $EMA_{20} = EMA_{20, t-1} \\times (1-\\alpha) + Close_t \\times \\alpha$\n"
        "- Seasonality: $Sin\\_DayOfWeek_t = \\sin(\\frac{2\\pi \\cdot DayOfWeek_t}{5})$\n"
        "- Noise: $Return\\_ZScore_t = \\frac{Return_t - \\mu_{20}}{\\sigma_{20} + 1e-8}$")

    add_paragraph_styled(doc, "2. Bảng kết quả so sánh ablation study đầy đủ 6 mô hình (Chương 4):", bold_prefix="Bảng Ablation Study: ")
    
    t_abl = doc.add_table(rows=1, cols=5)
    t_abl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_a = t_abl.rows[0].cells
    cols_a = ["Mô hình thực nghiệm", "Tập đặc trưng đầu vào", "MSE (USD)", "MAE (USD)", "Directional Accuracy (DA)"]
    for idx, c_name in enumerate(cols_a):
        hdr_a[idx].text = c_name
        set_cell_background(hdr_a[idx], "003366")
        set_cell_margins(hdr_a[idx], top=80, bottom=80, left=60, right=60)
        hdr_a[idx].paragraphs[0].runs[0].bold = True
        hdr_a[idx].paragraphs[0].runs[0].font.name = 'Times New Roman'
        hdr_a[idx].paragraphs[0].runs[0].font.size = Pt(9.5)

    abl_data = [
        ["Linear Regression", "Base temporal features", "73.9399", "4.5546", "53.40%"],
        ["LSTM", "Close-only temporal", "78.5781", "4.7485", "41.60%"],
        ["LSTM-GNN No-Gate", "Base node features", "75.1966", "4.6002", "42.20%"],
        ["LSTM-GNN Graph-Gate", "Base node features", "74.7984", "4.6045", "43.40%"],
        ["Graph-Gate + TSN Features", "Base + TSN features", "74.1900", "4.6156", "45.80%"],
        ["TSN-Attention Graph-Gated", "TSN + Temporal Attention + Noise", "73.7139", "4.5253", "48.80%"]
    ]
    for r_data in abl_data:
        row_cells = t_abl.add_row().cells
        for c_idx, val in enumerate(r_data):
            row_cells[c_idx].text = val
            set_cell_margins(row_cells[c_idx], top=50, bottom=50, left=50, right=50)
            row_cells[c_idx].paragraphs[0].runs[0].font.name = 'Times New Roman'
            row_cells[c_idx].paragraphs[0].runs[0].font.size = Pt(9)
            row_cells[c_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # IV. ĐIỀU CHỈNH CẤU TRÚC VÀ LỖI ĐÁNH SỐ
    add_heading_styled(doc, "IV. ĐIỀU CHỈNH CẤU TRÚC, ĐÁNH SỐ VÀ THAM CHIẾU HÌNH/BẢNG", level=1)
    
    struct_fixes = [
        ("Sửa lỗi đánh số Chương 5:", " Thay thế toàn bộ các tiểu mục '4.1. Ưu điểm', '4.2. Nhược điểm', '4.2. Hướng phát triển' thành '5.1. Ưu điểm', '5.2. Nhược điểm', '5.3. Hướng phát triển' để đúng quy chuẩn khoa học."),
        ("Bổ sung mục 1.3.2 và 2.3.2 bị thiếu:", " Tạo tiểu mục '1.3.2. Câu hỏi nghiên cứu' và '2.3.2. Các chỉ số đo lường lỗi phi tuyến' để lấp đầy khoảng trống cấu trúc."),
        ("Sửa trùng lặp 3.1.4.1:", " Sửa mục '3.1.4.1. Xây dựng cấu trúc...' thành '3.1.4.2. Xây dựng ma trận đồ thị động'."),
        ("Sửa lệch tham chiếu Hình/Bảng:", " Thay thế các tham chiếu sai lệch trong văn bản: sửa 'Hình 2' thành 'Hình 3' ở phần ma trận Pearson; sửa 'Bảng 5 và Hình 6' thành 'Bảng 4 và Hình 7' ở phần phân tích per-ticker."),
        ("Gộp các tiểu mục giải thích chênh lệch (4.3.1 - 4.3.6):", " Gộp 6 tiểu mục rườm rà thành 3 mục tinh gọn: '4.3.1. Khác biệt do mục tiêu đánh giá và cơ chế chuẩn hóa động', '4.3.2. Khác biệt do cấu trúc đồ thị động và batch training', '4.3.3. Khác biệt do xử lý streaming và kết luận thực tế'."),
        ("Rút gọn mục 4.2.4 (Đánh giá Graph-Gate):", " Cắt giảm 15 đoạn văn trùng lặp xuống còn 4 đoạn tập trung vào công thức cổng Gate và hệ số cổng Gate_Mean thực tế.")
    ]
    for idx, (f_title, f_desc) in enumerate(struct_fixes, 1):
        add_paragraph_styled(doc, f_desc, bold_prefix=f"{idx}. {f_title}")

    doc.add_paragraph()

    # V. BỔ SUNG TÀI LIỆU THAM KHẢO VÀ MINH BẠCH THAM SỐ
    add_heading_styled(doc, "V. BỔ SUNG TÀI LIỆU THAM KHẢO VÀ THAM SỐ THỰC TẾ", level=1)
    
    add_paragraph_styled(doc, 
        "Vị trí chèn: Chèn vào mục 'TÀI LIỆU THAM KHẢO' bị rỗng ở cuối báo cáo và phần Phương pháp Chương 3.")
    
    add_paragraph_styled(doc, 
        "- [1] D. Bahdanau, K. Cho, and Y. Bengio, 'Neural machine translation by jointly learning to align and translate', ICLR, 2015. (Bahdanau Attention)\n"
        "- [2] W. Hamilton, Z. Ying, and J. Leskovec, 'Inductive representation learning on large graphs', NeurIPS, 2017. (Graph Gate và GNN)\n"
        "- [3] R. J. Hyndman and G. Athanasopoulos, Forecasting: Principles and Practice, 3rd ed., OTexts, 2021. (Phân tách TSN)", 
        bold_prefix="1. Tài liệu tham khảo khoa học bổ sung: ")
        
    add_paragraph_styled(doc, 
        "Cần minh bạch các thông số thực tế đã được tinh chỉnh so với nghiên cứu gốc:\n"
        "- Lookback Window: Đã giảm từ 30 ngày (của nghiên cứu gốc) xuống còn 20 ngày ($LOOKBACK = 20$) để tối ưu thời gian huấn luyện và giảm thiểu độ nhiễu.\n"
        "- Chế độ Fast Mode: Để tối ưu hóa tốc độ huấn luyện cuốn chiếu trong Spark batch layer, hệ thống áp dụng $EXP\\_USE\\_FAST\\_MODE = True$ với 20 epochs huấn luyện ban đầu và 6 epochs cập nhật cho mỗi ngày tịnh tiến tiếp theo.", 
        bold_prefix="2. Minh bạch các tham số cấu hình: ")

    doc.save(doc_path)
    print("=" * 80)
    print(f"Recommendation document successfully generated at: {doc_path}")
    print("=" * 80)

if __name__ == "__main__":
    main()
